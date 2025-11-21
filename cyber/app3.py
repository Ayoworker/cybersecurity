"""
Cybersecurity Report Generator - Fixed Page Border Issue

Key Fix: Creates two separate documents (cover page and main report)
and merges them to ensure page border only appears on first page.
"""

import streamlit as st
import json
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io
import tempfile
import pandas as pd
from difflib import SequenceMatcher
import sqlite3
import hashlib

try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Set page config
st.set_page_config(
    page_title="Cybersecurity Report Generator",
    page_icon="🔒",
    layout="wide"
)

# Database Management
DB_PATH = "knowledge_base.db"

def init_database():
    """Initialize SQLite database for knowledge base - runs only once per session"""
    # Check if already initialized in this session
    if st.session_state.get('db_initialized'):
        return
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS knowledge_base (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            issue_name TEXT UNIQUE NOT NULL,
            implication TEXT NOT NULL,
            mitigation TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            usage_count INTEGER DEFAULT 0
        )
    ''')

    cursor.execute('''
        CREATE INDEX IF NOT EXISTS idx_issue_name ON knowledge_base(issue_name)
    ''')

    conn.commit()

    cursor.execute('SELECT COUNT(*) FROM knowledge_base')
    count = cursor.fetchone()[0]

    if count == 0:
        json_path = Path('knowledge_base.json')

        if json_path.exists():
            try:
                # Direct JSON loading without caching
                with open(json_path, 'r', encoding='utf-8') as f:
                    kb_data = json.load(f)

                imported = 0
                for issue_name, details in kb_data.items():
                    if isinstance(details, dict):
                        implication = details.get('implication', '')
                        mitigation = details.get('mitigation', '')
                        if implication and mitigation:
                            cursor.execute('''
                                INSERT OR IGNORE INTO knowledge_base (issue_name, implication, mitigation)
                                VALUES (?, ?, ?)
                            ''', (issue_name, implication, mitigation))
                            imported += 1

                conn.commit()
                st.success(f"✅ Imported {imported} entries from knowledge_base.json")

            except Exception as e:
                st.error(f"❌ Error loading knowledge_base.json: {e}")
        else:
            st.warning(f"⚠️ knowledge_base.json not found at {json_path.absolute()}")

    conn.close()
    st.session_state.db_initialized = True

@st.cache_data(ttl=3600)  # Cache for 1 hour
def load_kb_from_db():
    """Load all KB entries from database with caching"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT issue_name, implication, mitigation FROM knowledge_base ORDER BY usage_count DESC')
    rows = cursor.fetchall()
    conn.close()

    kb = {}
    for row in rows:
        kb[row[0]] = {
            'implication': row[1],
            'mitigation': row[2]
        }
    return kb

def add_to_kb_db(issue_name, implication, mitigation):
    """Add or update entry in database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    try:
        cursor.execute('''
            INSERT INTO knowledge_base (issue_name, implication, mitigation)
            VALUES (?, ?, ?)
            ON CONFLICT(issue_name) 
            DO UPDATE SET 
                implication=excluded.implication,
                mitigation=excluded.mitigation,
                updated_at=CURRENT_TIMESTAMP
        ''', (issue_name, implication, mitigation))
        conn.commit()
        success = True
    except Exception as e:
        st.error(f"Database error: {e}")
        success = False
    finally:
        conn.close()

    return success


def increment_kb_usage(issue_name):
    """Increment usage counter for KB entry with session state tracking"""
    # Use session state to track which issues we've already incremented in this session
    session_key = f'kb_used_{issue_name}'
    
    if not st.session_state.get(session_key):
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            UPDATE knowledge_base 
            SET usage_count = usage_count + 1 
            WHERE issue_name = ?
        ''', (issue_name,))
        conn.commit()
        conn.close()
        st.session_state[session_key] = True

def search_kb_db(query, top_n=5):
    """Search KB in database with similarity matching"""
    if not query or len(query) < 3:
        return []

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT issue_name, implication, mitigation, usage_count FROM knowledge_base')
    rows = cursor.fetchall()
    conn.close()

    query_lower = query.lower()
    matches = []

    for row in rows:
        issue_name, implication, mitigation, usage_count = row

        issue_similarity = SequenceMatcher(None, query_lower, issue_name.lower()).ratio()
        impl_similarity = SequenceMatcher(None, query_lower, implication.lower()).ratio()
        mitg_similarity = SequenceMatcher(None, query_lower, mitigation.lower()).ratio()

        max_similarity = max(issue_similarity, impl_similarity * 0.8, mitg_similarity * 0.8)

        keywords = query_lower.split()
        keyword_matches = sum(1 for keyword in keywords if len(keyword) > 3 and
                              (keyword in issue_name.lower() or
                               keyword in implication.lower() or
                               keyword in mitigation.lower()))

        if keyword_matches > 0:
            max_similarity += keyword_matches * 0.15

        usage_boost = min(usage_count * 0.01, 0.1)
        max_similarity += usage_boost

        if max_similarity > 0.2:
            matches.append({
                'issue': issue_name,
                'implication': implication,
                'mitigation': mitigation,
                'score': max_similarity,
                'usage_count': usage_count
            })

    matches.sort(key=lambda x: x['score'], reverse=True)
    return matches[:top_n]


def get_kb_stats():
    """Get KB statistics"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT COUNT(*), SUM(usage_count) FROM knowledge_base')
    total, total_usage = cursor.fetchone()
    conn.close()
    return {'total': total or 0, 'total_usage': total_usage or 0}


def export_kb_to_json():
    """Export entire KB to JSON"""
    kb = load_kb_from_db()
    return json.dumps(kb, indent=2)


def import_kb_from_json(json_data):
    """Import KB entries from JSON"""
    try:
        data = json.loads(json_data) if isinstance(json_data, str) else json_data
        count = 0
        for issue_name, details in data.items():
            if isinstance(details, dict):
                impl = details.get('implication', '')
                mitg = details.get('mitigation', '')
                if impl and mitg:
                    if add_to_kb_db(issue_name, impl, mitg):
                        count += 1
        return count
    except Exception as e:
        st.error(f"Import error: {e}")
        return 0


def extract_ips_from_word(file_bytes):
    """Extract IP addresses and hostnames from Word document tables"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        ip_list = []

        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]

                if row_idx == 0 and any(keyword in cells[0].lower() for keyword in ['ip', 'address', 'host']):
                    continue

                ip_pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}\b'

                for idx, cell_text in enumerate(cells):
                    ip_match = re.search(ip_pattern, cell_text)
                    if ip_match:
                        ip_addr = ip_match.group()
                        hostname = ''
                        if idx + 1 < len(cells):
                            hostname = cells[idx + 1]
                        elif idx > 0:
                            hostname = cells[idx - 1]

                        hostname = re.sub(ip_pattern, '', hostname).strip()
                        ip_list.append({'ip': ip_addr, 'host': hostname})
                        break

        return ip_list
    except Exception as e:
        st.error(f"Error reading Word document: {e}")
        return []


def extract_ips_from_pdf(file_bytes):
    """Extract IP addresses and hostnames from PDF"""
    if not PDF_AVAILABLE:
        st.error("PyPDF2 not installed. Run: pip install PyPDF2")
        return []

    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        ip_list = []
        ip_pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}\b'

        for page in pdf_reader.pages:
            text = page.extract_text()
            lines = text.split('\n')

            for line in lines:
                ip_match = re.search(ip_pattern, line)
                if ip_match:
                    ip_addr = ip_match.group()
                    hostname = line.replace(ip_addr, '').strip()
                    hostname = re.sub(r'^[\s\-\|:]+|[\s\-\|:]+$', '', hostname).strip()
                    ip_list.append({'ip': ip_addr, 'host': hostname})

        return ip_list
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return []


def extract_ips_from_csv(file_bytes):
    """Extract IP addresses from CSV file"""
    try:
        df = pd.read_csv(io.BytesIO(file_bytes))
        ip_list = []

        ip_col = None
        host_col = None

        for col in df.columns:
            col_lower = col.lower()
            if 'ip' in col_lower and ip_col is None:
                ip_col = col
            if any(keyword in col_lower for keyword in ['host', 'name', 'server']) and host_col is None:
                host_col = col

        if ip_col:
            for idx, row in df.iterrows():
                ip_addr = str(row[ip_col]).strip() if ip_col else ''
                hostname = str(row[host_col]).strip() if host_col and host_col in row else ''

                if ip_addr and ip_addr != 'nan':
                    ip_list.append({'ip': ip_addr, 'host': hostname})

        return ip_list
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        return []


def add_page_border_to_section(sectPr):
    """Add page border to a section properties element"""
    pgBorders = sectPr.find(qn('w:pgBorders'))

    if pgBorders is None:
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        sectPr.append(pgBorders)

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '48')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), 'FF0000')
        pgBorders.append(border)


def create_cover_page(doc: Document, app_name: str, version: str, author: str = None, logo_path: str = None):
    """Create cover page with red border and explicit section break"""

    # Add company logo at the top center (no spacing before it)
    if logo_path and Path(logo_path).exists():
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_path, width=Inches(3.0))
        except Exception:
            p = doc.add_paragraph()
            run = p.add_run("[Company Logo]")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run("[Company Logo]")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 4 spaces after logo
    for _ in range(4):
        doc.add_paragraph()

    # Blue text box with "Information Security Department"
    # Create a table to simulate a text box with blue background
    textbox_table = doc.add_table(rows=1, cols=1)
    textbox_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get the cell and add text
    cell = textbox_table.rows[0].cells[0]
    cell.text = "Information Security Department"

    # Center align the text
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set font size to 10
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            run.font.bold = True

    # Set blue background color for the cell
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '0000FF')
    cell._element.get_or_add_tcPr().append(shading_elm)

    # Set table width to be centered and reasonable size
    textbox_table.autofit = False
    textbox_table.allow_autofit = False
    for row in textbox_table.rows:
        for cell in row.cells:
            cell.width = Inches(4.5)

    # 3 spaces after text box
    for _ in range(3):
        doc.add_paragraph()

    # Application name and "Security Assessment"
    p_app = doc.add_paragraph()
    p_app.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_app = p_app.add_run(f"{app_name}\nSecurity Assessment")
    run_app.font.size = Pt(10)
    run_app.bold = True
    run_app.font.color.rgb = RGBColor(0, 0, 0)

    # 3 spaces after title
    for _ in range(3):
        doc.add_paragraph()

    # Version with current year
    current_year = datetime.now().year
    p_version = doc.add_paragraph()
    p_version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_version = p_version.add_run(f"Version 1.0 of {current_year}")
    run_version.font.size = Pt(10)

    # Date (on same line or next line - keeping original structure)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    current_date = datetime.now().strftime('%d %B %Y')
    run_date = p_date.add_run(current_date)
    run_date.font.size = Pt(10)

    # Spacing to push Document Control to bottom
    for _ in range(4):
        doc.add_paragraph()

    # Document Control at bottom left
    p_control = doc.add_paragraph()
    p_control.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_control = p_control.add_run("Document Control")
    run_control.font.size = Pt(10)
    run_control.bold = True

    # CRITICAL: Add section break with border for THIS section
    # The section properties go in the LAST paragraph of the section
    last_para = doc.add_paragraph()
    pPr = last_para._element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')

    # Add page size and margins - MODERATE MARGINS (0.75 inches)
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '12240')
    pgSz.set(qn('w:h'), '15840')
    sectPr.append(pgSz)

    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1080')  # 0.75 inch = 1080 twips
    pgMar.set(qn('w:right'), '1080')  # 0.75 inch
    pgMar.set(qn('w:bottom'), '1080')  # 0.75 inch
    pgMar.set(qn('w:left'), '1080')  # 0.75 inch
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)

    # Add RED BORDER to THIS section (the cover page)
    add_page_border_to_section(sectPr)

    # Add section type (next page)
    type_elem = OxmlElement('w:type')
    type_elem.set(qn('w:val'), 'nextPage')
    sectPr.append(type_elem)

    pPr.append(sectPr)


def add_document_info_page(doc: Document, app_name: str, author: str = None, logo_path: str = None):
    """Add page 2 with document information table"""
    # Remove the dynamic title - no longer needed

    # Document metadata table - 3 rows, 5 columns (to allow for splits)
    metadata_table = doc.add_table(rows=3, cols=5)
    metadata_table.style = 'Table Grid'

    # Row 1
    # Column 1: Logo (will be merged for full height)
    cell_logo = metadata_table.cell(0, 0)
    if logo_path and Path(logo_path).exists():
        try:
            # Clear any default text
            cell_logo.text = ""
            paragraph = cell_logo.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            # Scale logo to fit nicely in the cell
            run.add_picture(logo_path, width=Inches(1.5))
        except Exception:
            cell_logo.text = "[Logo]"
            cell_logo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        cell_logo.text = "[Logo]"
        cell_logo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Columns 2-3: App Name Security Assessment Review (left-aligned, merged)
    cell_title = metadata_table.cell(0, 1)
    metadata_table.cell(0, 1).merge(metadata_table.cell(0, 2))
    cell_title.text = f"{app_name}\nSecurity Assessment Review"
    cell_title.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in cell_title.paragraphs[0].runs:
        run.font.bold = True

    # Columns 4-5: FILE Reference (left-aligned, merged)
    cell_file_ref = metadata_table.cell(0, 3)
    metadata_table.cell(0, 3).merge(metadata_table.cell(0, 4))
    cell_file_ref.text = "FILE Reference:"
    cell_file_ref.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in cell_file_ref.paragraphs[0].runs:
        run.font.bold = True

    # Row 2
    # Column 1: Logo (will be merged)
    metadata_table.cell(1, 0).text = ""

    # Column 2: MODIFIED ON: (left part of split)
    cell_modified_label = metadata_table.cell(1, 1)
    cell_modified_label.text = "MODIFIED ON:"
    cell_modified_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in cell_modified_label.paragraphs[0].runs:
        run.font.bold = True

    # Column 3: [date] (right part of split, right aligned)
    current_date = datetime.now().strftime('%B %d, %Y')
    cell_modified_value = metadata_table.cell(1, 2)
    cell_modified_value.text = current_date
    cell_modified_value.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in cell_modified_value.paragraphs[0].runs:
        run.font.bold = True

    # Column 4: CONFIDENTIALITY: (left part of split)
    cell_conf_label = metadata_table.cell(1, 3)
    cell_conf_label.text = "CONFIDENTIALITY:"
    cell_conf_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in cell_conf_label.paragraphs[0].runs:
        run.font.bold = True

    # Column 5: CONFIDENTIAL (right part of split, right aligned)
    cell_conf_value = metadata_table.cell(1, 4)
    cell_conf_value.text = "CONFIDENTIAL"
    cell_conf_value.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in cell_conf_value.paragraphs[0].runs:
        run.font.bold = True

    # Row 3
    # Column 1: Logo (will be merged)
    metadata_table.cell(2, 0).text = ""

    # Column 2: VERSION: (left part of split)
    cell_version_label = metadata_table.cell(2, 1)
    cell_version_label.text = "VERSION:"
    cell_version_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in cell_version_label.paragraphs[0].runs:
        run.font.bold = True

    # Column 3: 1.0 OF [year] (right part of split, right aligned)
    current_year = datetime.now().year
    cell_version_value = metadata_table.cell(2, 2)
    cell_version_value.text = f"1.0 OF {current_year}"
    cell_version_value.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in cell_version_value.paragraphs[0].runs:
        run.font.bold = True

    # Column 4: Blank (split cell)
    metadata_table.cell(2, 3).text = ""

    # Column 5: Blank (split cell)
    metadata_table.cell(2, 4).text = ""

    # Merge column 1 cells (rows 0, 1, 2) to create full height logo column
    metadata_table.cell(0, 0).merge(metadata_table.cell(1, 0))
    metadata_table.cell(0, 0).merge(metadata_table.cell(2, 0))

    doc.add_paragraph()  # Spacing

    # Document details table
    details_table = doc.add_table(rows=4, cols=2)
    details_table.style = 'Table Grid'

    details_table.cell(0, 0).text = "File Name\n"
    details_table.cell(0, 1).text = f"{app_name} Security Assessment\n"

    details_table.cell(1, 0).text = "Compiled By:\n"
    details_table.cell(1, 1).text = "\n"  # Leave blank with spacing

    details_table.cell(2, 0).text = "Approved By:\n"
    details_table.cell(2, 1).text = "\n"

    details_table.cell(3, 0).text = "Submitted to:\n"
    details_table.cell(3, 1).text = "\n"

    # Make column 1 labels bold and add gray background
    for row in details_table.rows:
        # Bold text in column 1
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

        # Add gray background to column 1
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')  # Light gray
        row.cells[0]._element.get_or_add_tcPr().append(shading_elm)

    # Set column widths to match sign-off table proportions
    # Column 1: narrower (like "Position")
    for cell in details_table.columns[0].cells:
        cell.width = Inches(2.5)

    # Column 2: wider (like "Signature" + "Date" combined)
    for cell in details_table.columns[1].cells:
        cell.width = Inches(4.0)

    doc.add_paragraph()

    # Sign Off section
    p_signoff = doc.add_paragraph()
    run_signoff = p_signoff.add_run("Sign Off")
    run_signoff.font.bold = True
    run_signoff.font.size = Pt(14)

    # Sign off table
    signoff_table = doc.add_table(rows=3, cols=3)
    signoff_table.style = 'Table Grid'

    signoff_table.cell(0, 0).text = "Position\n"
    signoff_table.cell(0, 1).text = "Signature\n"
    signoff_table.cell(0, 2).text = "Date\n"

    # Make headers bold and add gray background to column 1
    for row_idx, row in enumerate(signoff_table.rows):
        if row_idx == 0:
            # Make all header cells bold
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

        # Add gray background to column 1 for all rows
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')  # Light gray
        row.cells[0]._element.get_or_add_tcPr().append(shading_elm)

    signoff_table.cell(1, 0).text = "\n"
    signoff_table.cell(1, 1).text = "\n"
    signoff_table.cell(1, 2).text = "\n"

    signoff_table.cell(2, 0).text = "\n"
    signoff_table.cell(2, 1).text = "\n"
    signoff_table.cell(2, 2).text = "\n"

    # Set column widths for sign-off table
    for cell in signoff_table.columns[0].cells:
        cell.width = Inches(2.5)
    for cell in signoff_table.columns[1].cells:
        cell.width = Inches(2.5)
    for cell in signoff_table.columns[2].cells:
        cell.width = Inches(1.5)

    doc.add_page_break()


def add_assessment_summary(doc: Document, app_name: str, system_arch_image: bytes = None):
    """Add assessment summary section"""
    # SUMMARY section
    p_summary_title = doc.add_paragraph()
    run_summary_title = p_summary_title.add_run('SUMMARY')
    run_summary_title.bold = True
    run_summary_title.underline = True
    run_summary_title.font.size = Pt(12)

    p_summary = doc.add_paragraph()
    p_summary.add_run(
        f'We performed a security assessment for {app_name} as detailed in the scope below. The security audit consisted of the following components:')

    # Add bulleted list with tick marks (✓)
    p_bullet1 = doc.add_paragraph()
    p_bullet1.add_run('✓  Servers and Network Infrastructure')
    p_bullet1.paragraph_format.left_indent = Inches(0.5)

    p_bullet2 = doc.add_paragraph()
    p_bullet2.add_run('✓  Architecture Review')
    p_bullet2.paragraph_format.left_indent = Inches(0.5)

    p_bullet3 = doc.add_paragraph()
    p_bullet3.add_run('✓  Mobile application Vulnerabilities')
    p_bullet3.paragraph_format.left_indent = Inches(0.5)

    p_bullet4 = doc.add_paragraph()
    p_bullet4.add_run('✓  Client attacks')
    p_bullet4.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()  # Spacing

    # SCOPE section
    p_scope_title = doc.add_paragraph()
    run_scope_title = p_scope_title.add_run('Scope')
    run_scope_title.bold = True
    run_scope_title.underline = True
    run_scope_title.font.size = Pt(12)

    p_scope = doc.add_paragraph()
    p_scope.add_run('The scope of this security vulnerability assessment was limited to the following.')

    doc.add_paragraph()  # Spacing

    # System Architecture (if provided)
    if system_arch_image:
        p_arch = doc.add_paragraph()
        p_arch.add_run('System Architecture:').bold = True

        doc.add_paragraph()
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                tmp.write(system_arch_image)
                tmp_path = tmp.name
            doc.add_picture(tmp_path, width=Inches(6))
            Path(tmp_path).unlink()
        except Exception:
            doc.add_paragraph('[Could not insert architecture diagram]')

    doc.add_page_break()


def add_ip_inventory_table(doc: Document, ip_inventory: list):
    """Add IP inventory table to document"""
    doc.add_heading('IP Address Inventory', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'IP Address'
    table.cell(0, 1).text = 'Host Name'
    for ip in ip_inventory:
        row = table.add_row().cells
        row[0].text = ip.get('ip', '')
        row[1].text = ip.get('host', '')
    doc.add_paragraph()


def add_findings_master_table(doc: Document, findings: list):
    """Add findings master list table grouped by classification"""
    doc.add_heading('Findings Master List', level=2)
    # Start with 5 columns for the main table structure
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Header row - merge cells to show only 4 visible columns
    hdr = table.rows[0].cells
    hdr[0].text = 'No.'
    hdr[1].text = 'Issue'

    # Merge severity and status cells for header to show as one column
    hdr[2].merge(hdr[3])
    hdr[2].text = 'Severity'  # This spans 2 columns but appears as one

    hdr[4].text = 'Responsibility'

    #table.autofit = False
    #table.allow_autofit = False  # depending on python-docx version

    # Make header row bold
    for cell in [hdr[0], hdr[1], hdr[2], hdr[4]]:  # Skip the merged cell that doesn't exist
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

            # Set column widths in centimeters (converted to inches)
            # 1 cm = 0.3937 inches
    for cell in table.columns[0].cells:
        cell.width = Inches(1.5 * 0.3937)  # 1.5 cm - No. column
    for cell in table.columns[1].cells:
        cell.width = Inches(12.0 * 0.3937)  # 12.0 cm - Issue column (widest)
    for cell in table.columns[2].cells:
        cell.width = Inches(3.0 * 0.3937)  # 3.0 cm - Severity Level part
    for cell in table.columns[3].cells:
        cell.width = Inches(3.0 * 0.3937)  # 3.0 cm - Status part
    for cell in table.columns[4].cells:
        cell.width = Inches(3.0 * 0.3937)  # 3.0 cm - Responsibility

    # Group findings by classification
    classifications = ['Mobile Application Vulnerability', 'Server Vulnerabilities', 'Web Vulnerabilities']

    for classification in classifications:
        # Filter findings for this classification
        classified_findings = [f for f in findings if f.get('classification') == classification]

        if classified_findings:
            # Add classification header row - spans all 5 columns
            row = table.add_row().cells
            row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])
            row[0].text = classification

            # Make classification row bold
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

            # Add findings for this classification - each has 5 separate columns
            for f in classified_findings:
                row = table.add_row().cells
                row[0].text = f.get('number', '')
                row[1].text = f.get('issue', '')

                # Split severity into level and status in separate columns
                severity_level = f.get('severity_level', 'Medium')
                severity_status = f.get('severity_status', 'Open')

                row[2].text = severity_level  # Severity Level
                row[3].text = severity_status  # Status

                # Apply background color based on status to the status column only
                status_shading = OxmlElement('w:shd')
                if severity_status.lower() == 'open':
                    status_shading.set(qn('w:fill'), 'FF0000')  # Red
                elif severity_status.lower() == 'closed':
                    status_shading.set(qn('w:fill'), '00FF00')  # Green
                row[3]._element.get_or_add_tcPr().append(status_shading)

                row[4].text = f.get('responsible_party', '')

    doc.add_page_break()

def resize_image_for_table(img_bytes, max_width=1200):
    """Resize image to fit nicely in table cell while maintaining aspect ratio"""
    try:
        img = Image.open(io.BytesIO(img_bytes))

        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background

        if img.width > max_width:
            ratio = max_width / img.width
            new_height = int(img.height * ratio)
            img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)

        output = io.BytesIO()
        img.save(output, format='JPEG', quality=85, optimize=True)
        output.seek(0)
        return output.read()
    except Exception:
        return img_bytes


def set_column_width(column, width_cm):
    """Set column width in cm (converted to inches)"""
    for cell in column.cells:
        cell.width = Inches(width_cm / 2.54)


def generate_finding_pages(doc: Document, findings: list, ip_inventory: list, uploaded_images: dict):
    """Generate detailed finding pages"""
    kb = load_kb_from_db()

    for f in findings:
        finding_num = f.get('number', '')
        issue = f.get('issue', '')
        severity = f.get('severity', '')
        responsible = f.get('responsible_party', '')

        custom_implication = f.get('implication', '')
        custom_mitigation = f.get('mitigation', '')

        affected_hosts = f.get('affected_hosts', [])
        if isinstance(affected_hosts, list):
            hosts_str = ', '.join(affected_hosts) if affected_hosts else '[No hosts specified]'
        else:
            hosts_str = '[No hosts specified]'

        doc.add_heading(f"Finding {finding_num}: {issue}", level=2)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'

        # Set column widths: 2.11 cm for first column, 18 cm for second column
        set_column_width(table.columns[0], 2.11)
        set_column_width(table.columns[1], 18)

        implication = custom_implication
        mitigation = custom_mitigation

        if not implication or not mitigation:
            kb_key = None
            if issue in kb:
                kb_key = issue
                increment_kb_usage(issue)
            else:
                for k in kb.keys():
                    if k.lower() in issue.lower() or issue.lower() in k.lower():
                        kb_key = k
                        increment_kb_usage(k)
                        break

            if kb_key:
                if not implication:
                    implication = kb.get(kb_key, {}).get('implication', '[No implication provided]')
                if not mitigation:
                    mitigation = kb.get(kb_key, {}).get('mitigation', '[No mitigation provided]')

        if not implication:
            implication = '[No implication provided]'
        if not mitigation:
            mitigation = '[No mitigation provided]'

        rows_data = [
            (f"Finding {finding_num}", issue),
            ("Affected Host", hosts_str),
            ("Implication", implication),
            ("Risk Rating", severity),
            ("Detail", None),
            ("Mitigation", mitigation),
            ("Comment", '')
        ]

        for i, (left, right) in enumerate(rows_data):
            table.cell(i, 0).text = left
            # Add gray background to first column
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D3D3D3')  # Gray background
            table.cell(i, 0)._element.get_or_add_tcPr().append(shading_elm)

            if right is not None:
                table.cell(i, 1).text = right

        detail_cell = table.cell(4, 1)
        finding_images = uploaded_images.get(finding_num, [])

        if finding_images:
            detail_cell.text = ''

            for img_idx, img_bytes in enumerate(finding_images):
                try:
                    resized_img_bytes = resize_image_for_table(img_bytes, max_width=1200)

                    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp:
                        tmp.write(resized_img_bytes)
                        tmp_path = tmp.name

                    if img_idx == 0:
                        paragraph = detail_cell.paragraphs[0]
                    else:
                        paragraph = detail_cell.add_paragraph()

                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0

                    run = paragraph.add_run()
                    run.add_picture(tmp_path, width=Inches(4.0))

                    Path(tmp_path).unlink()
                except Exception as e:
                    detail_cell.text = f'[Could not insert evidence image {img_idx + 1}: {str(e)}]'
        else:
            detail_cell.text = '[No evidence image provided for this finding]'

        doc.add_page_break()


def generate_report(data: dict, uploaded_images: dict, arch_image: bytes = None, logo_image: bytes = None):
    """Generate the Word document report with proper section breaks"""
    app_name = data.get('application_name', 'Application')
    author = data.get('author')
    version = "v1.0"

    # Save logo to temp file if provided
    logo_path = None
    if logo_image:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                tmp.write(logo_image)
                logo_path = tmp.name
        except Exception:
            logo_path = None

    # Create single document with all content
    doc = Document()

    # Create cover page with border and section break
    create_cover_page(doc, app_name, version, author, logo_path)

    # Clean up logo temp file
    if logo_path:
        try:
            Path(logo_path).unlink()
        except Exception:
            pass

    # Add remaining content (all in the new section WITHOUT borders)
    add_document_info_page(doc, app_name, author, logo_path)
    add_assessment_summary(doc, app_name, arch_image)

    ip_inventory = data.get('ip_inventory', [])
    findings = data.get('findings', [])

    add_ip_inventory_table(doc, ip_inventory)
    add_findings_master_table(doc, findings)
    generate_finding_pages(doc, findings, ip_inventory, uploaded_images)

    # Save to BytesIO
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    safe_name = re.sub(r"[^0-9a-zA-Z-_ ]+", '', app_name).strip().replace(' ', '_')
    filename = f"{safe_name}_{version}.docx"

    return bio, filename


def main():
    """Main Streamlit application"""
    # Initialize session state for database
    if 'db_initialized' not in st.session_state:
        st.session_state.db_initialized = False
        
    # Initialize database (will only run once per session)
    init_database()


    st.title("🔒 Cybersecurity Report Generator")
    st.markdown("Generate professional cybersecurity assessment reports - No technical knowledge required!")

     # New Report button at the top
    if st.button("🆕 New Report", type="secondary", help="Reset all inputs and start a new report"):
        # Clear all session state
        st.session_state.ip_inventory = [{'ip': '', 'host': ''}]
        st.session_state.findings = []
        st.session_state.images = {}
        st.session_state.arch_image = None
        st.session_state.mobile_arch_image = None
        st.session_state.logo_image = None
        st.session_state.file_processed = False
        st.session_state.scroll_position = 0
        if 'new_finding_idx' in st.session_state:
            del st.session_state['new_finding_idx']
        st.success("✅ All inputs cleared! Ready for a new report.")
        st.rerun()

    # Add JavaScript to handle scroll position after new finding is added
    if 'new_finding_idx' in st.session_state:
        new_idx = st.session_state['new_finding_idx']
        scroll_js = f"""
        <script>
            setTimeout(function() {{
                var element = document.getElementById('finding_{new_idx}');
                if (element) {{
                    element.scrollIntoView({{ behavior: 'smooth', block: 'center' }});
                }}
            }}, 100);
        </script>
        """
        st.markdown(scroll_js, unsafe_allow_html=True)
        del st.session_state['new_finding_idx']

    # Initialize session state
    if 'ip_inventory' not in st.session_state:
        st.session_state.ip_inventory = [{'ip': '', 'host': ''}]
    if 'findings' not in st.session_state:
        st.session_state.findings = []
    if 'images' not in st.session_state:
        st.session_state.images = {}
    if 'arch_image' not in st.session_state:
        st.session_state.arch_image = None
    if 'mobile_arch_image' not in st.session_state:
        st.session_state.mobile_arch_image = None
    if 'logo_image' not in st.session_state:
        st.session_state.logo_image = None
    if 'file_processed' not in st.session_state:
        st.session_state.file_processed = False
    if 'scroll_position' not in st.session_state:
        st.session_state.scroll_position = 0

    # Sidebar for app info and KB
    with st.sidebar:
        st.header("📋 Report Information")
        app_name = st.text_input("Application Name*", placeholder="e.g., MyApp Security Assessment")
        author = st.text_input("Author Name", placeholder="e.g., Security Team")

        st.divider()

        st.header("📸 Optional Uploads")

        # Company Logo for Cover Page
        with st.expander("Company Logo (Cover Page)"):
            logo_file = st.file_uploader("Upload company logo", type=['png', 'jpg', 'jpeg'], key='logo_upload')
            if logo_file:
                st.session_state.logo_image = logo_file.read()
                st.success("✅ Company Logo Uploaded")
                st.image(st.session_state.logo_image, width=200, caption="Logo Preview")

        # System Architecture Diagram
        with st.expander("System Architecture Diagram"):
            arch_file = st.file_uploader("Upload system diagram", type=['png', 'jpg', 'jpeg'], key='arch_upload')
            if arch_file:
                st.session_state.arch_image = arch_file.read()
                st.success("✅ System Architecture Uploaded")

        # Mobile Application Architecture
        with st.expander("Mobile Application Architecture"):
            mobile_arch_file = st.file_uploader("Upload mobile architecture diagram", type=['png', 'jpg', 'jpeg'],
                                                key='mobile_arch_upload')
            if mobile_arch_file:
                st.session_state.mobile_arch_image = mobile_arch_file.read()
                st.success("✅ Mobile Architecture Uploaded")

        with st.expander("📚 Knowledge Base Statistics"):
            stats = get_kb_stats()
            col_stat1, col_stat2 = st.columns(2)
            with col_stat1:
                st.metric("Total KB Entries", stats['total'])
            with col_stat2:
                st.metric("Total Usage", stats['total_usage'])

            st.caption(
                "The knowledge base is automatically loaded from knowledge_base.json on first run and stored in SQLite.")

            st.divider()

            st.markdown("**Export Knowledge Base**")
            if st.button("📥 Export KB to JSON", use_container_width=True):
                kb_json = export_kb_to_json()
                st.download_button(
                    label="💾 Download JSON File",
                    data=kb_json,
                    file_name=f"kb_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True
                )

    # Main content area
    st.header("📊 Step 1: IP Address Inventory")
    st.markdown("Add IP addresses and their corresponding host names")

    with st.expander("📤 Quick Import: Upload Document with IP Addresses"):
        st.markdown("""
        Upload a document containing IP addresses and hostnames. The app will automatically extract them.

        **Supported formats:**
        - Word (.docx) - Best for tables
        - PDF (.pdf) - Extracts text with IPs
        - CSV (.csv) - Must have IP and hostname columns
        """)

        upload_col1, upload_col2 = st.columns([3, 1])

        with upload_col1:
            ip_file = st.file_uploader(
                "Upload document",
                type=['docx', 'pdf', 'csv'],
                key='ip_file_upload',
                help="Upload a Word, PDF, or CSV file containing IP addresses"
            )

        with upload_col2:
            st.write("")
            st.write("")
            if st.button("🔄 Clear All IPs", help="Remove all IP entries and start fresh"):
                st.session_state.ip_inventory = [{'ip': '', 'host': ''}]
                st.session_state.file_processed = False
                st.rerun()

        if st.session_state.get('file_processed', False):
            st.info("✅ File processed! IPs loaded below. You can now edit them or add more.")
            st.session_state.file_processed = False

        if ip_file:
            if st.button("📥 Import IPs from File", type="primary", use_container_width=True):
                file_type = ip_file.name.split('.')[-1].lower()
                file_bytes = ip_file.read()

                with st.spinner(f"Extracting IPs from {file_type.upper()}..."):
                    extracted_ips = []

                    if file_type == 'docx':
                        extracted_ips = extract_ips_from_word(file_bytes)
                    elif file_type == 'pdf':
                        extracted_ips = extract_ips_from_pdf(file_bytes)
                    elif file_type == 'csv':
                        extracted_ips = extract_ips_from_csv(file_bytes)

                    if extracted_ips:
                        seen = set()
                        unique_ips = []
                        for ip_entry in extracted_ips:
                            ip_key = ip_entry['ip']
                            if ip_key not in seen:
                                seen.add(ip_key)
                                unique_ips.append(ip_entry)

                        st.session_state.ip_inventory = unique_ips
                        st.session_state['file_processed'] = True
                        st.success(f"✅ Extracted {len(unique_ips)} unique IP addresses!")
                        st.rerun()
                    else:
                        st.warning("⚠️ No IP addresses found in the document. Please check the format.")

    st.subheader("IP Addresses and Hostnames")

    for idx, ip_entry in enumerate(st.session_state.ip_inventory):
        col1, col2, col3 = st.columns([3, 3, 1])
        with col1:
            ip = st.text_input(
                f"IP Address",
                value=ip_entry.get('ip', ''),
                key=f'ip_{idx}',
                placeholder="e.g., 192.168.1.10",
                label_visibility="collapsed"
            )
            st.session_state.ip_inventory[idx]['ip'] = ip
        with col2:
            host = st.text_input(
                f"Host Name",
                value=ip_entry.get('host', ''),
                key=f'host_{idx}',
                placeholder="e.g., web-server-01",
                label_visibility="collapsed"
            )
            st.session_state.ip_inventory[idx]['host'] = host
        with col3:
            st.write("")
            st.write("")
            if st.button("🗑️", key=f'del_ip_{idx}', help="Delete this entry"):
                st.session_state.ip_inventory.pop(idx)
                st.rerun()

    if st.button("➕ Add IP Address", type="secondary"):
        st.session_state.ip_inventory.append({'ip': '', 'host': ''})
        st.rerun()

    if any(ip.get('ip') or ip.get('host') for ip in st.session_state.ip_inventory):
        with st.expander("📋 Preview IP Inventory"):
            preview_df = pd.DataFrame([
                {'IP Address': ip.get('ip', ''), 'Host Name': ip.get('host', '')}
                for ip in st.session_state.ip_inventory
                if ip.get('ip') or ip.get('host')
            ])
            st.dataframe(preview_df, use_container_width=True, hide_index=True)

    st.divider()

    # Findings Section
    st.header("🔍 Step 2: Security Findings")
    st.markdown("Add all security findings discovered during the assessment")

    if not st.session_state.findings:
        st.info("👇 Click 'Add Finding' button below to start adding security findings")
    else:
        for idx, finding in enumerate(st.session_state.findings):
            # Create anchor for this finding
            finding_anchor = f"finding_{idx}"
            st.markdown(f'<div id="{finding_anchor}"></div>', unsafe_allow_html=True)

            with st.expander(
                    f"**Finding {finding.get('number', idx + 1)}** - {finding.get('issue', 'New Finding')[:50]}...",
                    expanded=idx == len(st.session_state.findings) - 1):
                col1, col2 = st.columns([1, 3])

                with col1:
                    number = st.text_input(
                        "Finding Number*",
                        value=finding.get('number', ''),
                        key=f'find_num_{idx}',
                        placeholder="e.g., 1, 1.1, 2.3"
                    )

                with col2:
                    issue = st.text_area(
                        "Issue Description*",
                        value=finding.get('issue', ''),
                        key=f'issue_{idx}',
                        placeholder="Describe the security issue found...",
                        height=100
                    )

                    # Show KB suggestions for issue description as a dropdown
                    if issue and len(issue) > 3:
                        similar_entries = search_kb_db(issue)
                        if similar_entries:
                            # Create options for selectbox with truncated preview
                            suggestion_options = ["-- Select a similar issue from KB --"]
                            for entry in similar_entries:
                                # Truncate issue name to one line
                                issue_preview = entry['issue'][:70] + "..." if len(entry['issue']) > 70 else entry[
                                    'issue']
                                # Remove newlines
                                issue_preview = issue_preview.replace('\n', ' ').replace('\r', '')
                                suggestion_options.append(issue_preview)

                            selected_issue = st.selectbox(
                                "💡 Similar issues in Knowledge Base:",
                                options=suggestion_options,
                                key=f'issue_select_{idx}',
                                help="Select to auto-fill issue, implication, and mitigation"
                            )

                            # If user selected a suggestion, update all fields
                            if selected_issue != "-- Select a similar issue from KB --":
                                selected_idx = suggestion_options.index(selected_issue) - 1
                                st.session_state.findings[idx]['issue'] = similar_entries[selected_idx]['issue']
                                st.session_state.findings[idx]['implication'] = similar_entries[selected_idx][
                                    'implication']
                                st.session_state.findings[idx]['mitigation'] = similar_entries[selected_idx][
                                    'mitigation']
                                increment_kb_usage(similar_entries[selected_idx]['issue'])
                                st.rerun()

                # Classification dropdown (full width below issue description)
                classification = st.selectbox(
                    "Classification*",
                    ['Mobile Application Vulnerability', 'Server Vulnerabilities', 'Web Vulnerabilities'],
                    index=['Mobile Application Vulnerability', 'Server Vulnerabilities', 'Web Vulnerabilities'].index(
                        finding.get('classification', 'Mobile Application Vulnerability')
                    ),
                    key=f'classification_{idx}',
                    help="Select the vulnerability classification"
                )

                col3, col4, col5 = st.columns(3)

                with col3:
                    severity_level = st.selectbox(
                        "Severity Level*",
                        ['Low', 'Medium', 'High', 'Critical'],
                        index=['Low', 'Medium', 'High', 'Critical'].index(finding.get('severity_level', 'Medium')),
                        key=f'sev_level_{idx}'
                    )

                with col4:
                    severity_status = st.selectbox(
                        "Status*",
                        ['Open', 'Closed'],
                        index=['Open', 'Closed'].index(finding.get('severity_status', 'Open')),
                        key=f'sev_status_{idx}'
                    )

                with col5:
                    responsible = st.text_input(
                        "Responsible Party*",
                        value=finding.get('responsible_party', ''),
                        key=f'resp_{idx}',
                        placeholder="e.g., IT Team, Vendor"
                    )

                st.divider()

                st.markdown("**Affected Hosts** (Select from IP inventory)")

                # Get IP addresses from inventory
                available_ips = [ip.get('ip', '') for ip in st.session_state.ip_inventory if ip.get('ip')]

                if available_ips:
                    current_hosts = finding.get('affected_hosts', [])
                    if not isinstance(current_hosts, list):
                        current_hosts = []

                    selected_hosts = st.multiselect(
                        "Select affected IP addresses:",
                        options=available_ips,
                        default=current_hosts,
                        key=f'hosts_{idx}',
                        help="Select one or more IP addresses affected by this finding"
                    )

                    st.session_state.findings[idx]['affected_hosts'] = selected_hosts
                else:
                    st.warning("⚠️ No IP addresses available. Please add IPs in Step 1.")
                    st.session_state.findings[idx]['affected_hosts'] = []

                st.divider()

                st.markdown("**Implication** (What could happen if not fixed)")

                # Auto-fill implication from KB if issue matches
                if issue and not finding.get('implication'):
                    kb = load_kb_from_db()
                    kb_match = None

                    if issue in kb:
                        kb_match = issue
                    else:
                        for k in kb.keys():
                            if k.lower() in issue.lower() or issue.lower() in k.lower():
                                kb_match = k
                                break

                    if kb_match:
                        st.session_state.findings[idx]['implication'] = kb.get(kb_match, {}).get('implication', '')
                        increment_kb_usage(kb_match)

                implication = st.text_area(
                    "Describe the security implication",
                    value=finding.get('implication', ''),
                    key=f'impl_{idx}',
                    placeholder="This will auto-fill if issue is found in Knowledge Base...",
                    height=100
                )

                st.divider()

                st.markdown("**Mitigation** (How to fix this issue)")

                # Auto-fill mitigation from KB if issue matches
                if issue and not finding.get('mitigation'):
                    kb = load_kb_from_db()
                    kb_match = None

                    if issue in kb:
                        kb_match = issue
                    else:
                        for k in kb.keys():
                            if k.lower() in issue.lower() or issue.lower() in k.lower():
                                kb_match = k
                                break

                    if kb_match:
                        st.session_state.findings[idx]['mitigation'] = kb.get(kb_match, {}).get('mitigation', '')
                        increment_kb_usage(kb_match)

                mitigation = st.text_area(
                    "Describe the mitigation steps",
                    value=finding.get('mitigation', ''),
                    key=f'mitg_{idx}',
                    placeholder="This will auto-fill if issue is found in Knowledge Base...",
                    height=100
                )

                if issue and implication and mitigation:
                    kb = load_kb_from_db()
                    if issue not in kb:
                        st.info(f"💾 This issue is not in the Knowledge Base yet")
                        if st.button(f"💾 Save '{issue[:30]}...' to Knowledge Base", key=f"save_kb_{idx}",
                                     type="secondary"):
                            if add_to_kb_db(issue, implication, mitigation):
                                st.success(f"✅ Added to Knowledge Base!")
                    else:
                        st.success("✅ This issue is already in the Knowledge Base")

                st.divider()

                st.markdown("**Evidence Images (Optional)**")
                st.caption("Upload one or more images. They will be stacked vertically in the report.")

                current_images = st.session_state.images.get(number, [])
                if current_images:
                    st.success(f"✅ {len(current_images)} image(s) uploaded")
                    cols = st.columns(min(len(current_images), 3))
                    for img_idx, img_bytes in enumerate(current_images):
                        with cols[img_idx % 3]:
                            st.image(img_bytes, caption=f"Image {img_idx + 1}", width=150)
                            if st.button(f"🗑️ Remove", key=f'remove_img_{idx}_{img_idx}'):
                                st.session_state.images[number].pop(img_idx)
                                st.rerun()

                img_files = st.file_uploader(
                    f"Select evidence images for Finding {number}",
                    type=['png', 'jpg', 'jpeg'],
                    key=f'img_uploader_{idx}',
                    label_visibility="collapsed",
                    accept_multiple_files=True
                )

                if img_files and st.button(f"📤 Add Selected Images", key=f'add_img_{idx}', type="secondary"):
                    if number not in st.session_state.images:
                        st.session_state.images[number] = []

                    added_count = 0
                    for img_file in img_files:
                        img_bytes = img_file.read()
                        img_hash = hashlib.md5(img_bytes).hexdigest()
                        existing_hashes = [hashlib.md5(existing_img).hexdigest()
                                           for existing_img in st.session_state.images[number]]

                        if img_hash not in existing_hashes:
                            try:
                                resized_img_bytes = resize_image_for_table(img_bytes)
                                st.session_state.images[number].append(resized_img_bytes)
                                added_count += 1
                            except Exception as e:
                                st.warning(f"Could not process image: {e}")
                                st.session_state.images[number].append(img_bytes)
                                added_count += 1

                    if added_count > 0:
                        st.success(f"✅ Added {added_count} new image(s)")
                        st.rerun()
                    else:
                        st.info("ℹ️ No new images added (duplicates detected)")

                combined_severity = f"{severity_level} - {severity_status}"

                st.session_state.findings[idx] = {
                    'number': number,
                    'issue': issue,
                    'classification': classification,
                    'severity_level': severity_level,
                    'severity_status': severity_status,
                    'severity': combined_severity,
                    'responsible_party': responsible,
                    'implication': implication,
                    'mitigation': mitigation,
                    'affected_hosts': st.session_state.findings[idx].get('affected_hosts', [])
                }

                col_del1, col_del2 = st.columns([4, 1])
                with col_del2:
                    if st.button("🗑️ Delete Finding", key=f'del_find_{idx}', type="secondary",
                                 use_container_width=True):
                        st.session_state.findings.pop(idx)
                        if number in st.session_state.images:
                            del st.session_state.images[number]
                        st.rerun()

    st.write("")

    add_button_anchor = "add_finding_section"
    st.markdown(f'<div id="{add_button_anchor}"></div>', unsafe_allow_html=True)

    if st.button("➕ Add New Finding", type="primary", use_container_width=True, key="add_finding_bottom"):
        new_number = str(len(st.session_state.findings) + 1)
        new_idx = len(st.session_state.findings)
        st.session_state.findings.append({
            'number': new_number,
            'issue': '',
            'severity_level': 'Medium',
            'severity_status': 'Open',
            'responsible_party': '',
            'implication': '',
            'mitigation': '',
            'affected_hosts': []
        })
        st.session_state['new_finding_idx'] = new_idx
        st.rerun()

    if st.session_state.findings:
        st.divider()
        st.subheader("📊 Findings Summary")
        summary_data = []
        for f in st.session_state.findings:
            if f.get('number') or f.get('issue'):
                summary_data.append({
                    'No.': f.get('number', ''),
                    'Issue': f.get('issue', '')[:50] + ('...' if len(f.get('issue', '')) > 50 else ''),
                    'Severity': f.get('severity_level', ''),
                    'Status': f.get('severity_status', ''),
                    'Responsible': f.get('responsible_party', '')
                })

        if summary_data:
            df = pd.DataFrame(summary_data)
            st.dataframe(df, use_container_width=True, hide_index=True)

    st.divider()

    # Generate Report Section
    st.header("📄 Step 3: Generate Report")

    can_generate = True
    issues = []

    if not app_name:
        issues.append("❌ Application Name is required")
        can_generate = False

    if not st.session_state.findings:
        issues.append("❌ At least one finding is required")
        can_generate = False
    else:
        for idx, f in enumerate(st.session_state.findings):
            if not f.get('number'):
                issues.append(f"❌ Finding {idx + 1} is missing a number")
                can_generate = False
            if not f.get('issue'):
                issues.append(f"❌ Finding {idx + 1} is missing an issue description")
                can_generate = False

    if issues:
        st.warning("Please fix the following issues before generating:")
        for issue in issues:
            st.markdown(issue)
    else:
        st.success("✅ Ready to generate report!")

    if st.button("🚀 Generate Report", type="primary", use_container_width=True, disabled=not can_generate):
        with st.spinner("Generating your cybersecurity report..."):
            try:
                report_data = {
                    'application_name': app_name,
                    'author': author,
                    'ip_inventory': [ip for ip in st.session_state.ip_inventory if ip.get('ip') or ip.get('host')],
                    'findings': st.session_state.findings
                }

                report_bytes, filename = generate_report(
                    report_data,
                    st.session_state.images,
                    st.session_state.arch_image,
                    st.session_state.logo_image
                )

                st.success("✅ Report generated successfully!")

                st.download_button(
                    label="📥 Download Report (.docx)",
                    data=report_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary"
                )

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Total Findings", len(st.session_state.findings))
                with col2:
                    st.metric("IP Addresses", len([ip for ip in st.session_state.ip_inventory if ip.get('ip')]))
                with col3:
                    st.metric("Evidence Images", sum(len(imgs) for imgs in st.session_state.images.values()))

            except Exception as e:
                st.error(f"❌ Error generating report: {e}")
                st.exception(e)


if __name__ == "__main__":
    main()



